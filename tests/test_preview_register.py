"""Tests for development-only loading of the approved Word register."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from app.config import Settings
from app.main import ReadinessChecks, create_app
from app.preview_register import load_preview_register


HEADERS = (
    "Applicant",
    "Degree",
    "Age",
    "Academic age (years)",
    "Gender",
    "First-author papers",
    "Last-author papers",
    "Total papers",
    "h-index",
    "Total citations",
    "ORCID",
    "Google Scholar citations",
    "GS identity certainty",
)


def test_load_preview_register_preserves_blanks_and_numeric_values(tmp_path: Path) -> None:
    source = tmp_path / "register.docx"
    document = Document()
    table = document.add_table(rows=2, cols=len(HEADERS))
    for cell, value in zip(table.rows[0].cells, HEADERS):
        cell.text = value
    values = (
        "Applicant One", "PhD", "35", "7.5", "", "6", "2", "14", "11", "420",
        "0000-0002-1825-0097", "500", "High",
    )
    for cell, value in zip(table.rows[1].cells, values):
        cell.text = value
    document.save(source)

    records = load_preview_register(source)

    assert len(records) == 1
    assert records[0].applicant == "Applicant One"
    assert records[0].age == 35
    assert records[0].academic_age == 7.5
    assert records[0].gender is None
    assert records[0].total_citations == 420


def test_load_preview_register_refuses_wrong_header_or_missing_file(tmp_path: Path) -> None:
    assert load_preview_register(tmp_path / "missing.docx") == ()
    source = tmp_path / "wrong.docx"
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Wrong"
    document.save(source)

    assert load_preview_register(source) == ()


def test_real_data_preview_requires_explicit_gate_and_loopback_host(
    tmp_path: Path, monkeypatch
) -> None:
    source = tmp_path / "register.docx"
    document = Document()
    table = document.add_table(rows=2, cols=len(HEADERS))
    for cell, value in zip(table.rows[0].cells, HEADERS):
        cell.text = value
    values = (
        "Applicant One", "PhD", "35", "7.5", "", "6", "2", "14", "11", "420",
        "0000-0002-1825-0097", "500", "High",
    )
    for cell, value in zip(table.rows[1].cells, values):
        cell.text = value
    document.save(source)
    monkeypatch.setenv("EHF_PREVIEW_REAL_DATA_ENABLED", "true")
    monkeypatch.setenv("EHF_PREVIEW_REGISTER_PATH", str(source))
    app = create_app(
        Settings.from_environment({"EHF_ALLOWED_HOST": "localhost"}),
        readiness_checks=ReadinessChecks(lambda _: None, lambda _: None),
    )

    local = TestClient(app, base_url="http://localhost").get(
        "/__preview/internal/administrator/"
    )
    public = TestClient(app, base_url="http://ehf.isab.science").get(
        "/__preview/internal/administrator/"
    )

    assert local.status_code == 200
    assert "Applicant One" in local.text
    assert public.status_code != 200
    assert "Applicant One" not in public.text
