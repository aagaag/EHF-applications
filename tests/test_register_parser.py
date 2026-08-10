from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.importer.register import RegisterParseError, parse_register


HEADERS = (
    "Applicant",
    "Degree",
    "Age",
    "Academic age (years)",
    "Gender",
    "First-author papers",
    "Last-author papers",
    "Total papers",
    "h-index",
    "Total citations",
    "ORCID",
    "Google Scholar citations",
    "GS identity certainty",
)


def write_register(
    path: Path,
    *,
    first_age: str = "31",
    first_academic_age: str = "4.8",
    scholar_citations: str = "",
    total_citations: str = "5",
) -> None:
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Unrelated table"
    table = document.add_table(rows=1, cols=len(HEADERS))
    for cell, header in zip(table.rows[0].cells, HEADERS, strict=True):
        cell.text = header.upper()
    for number in range(36):
        cells = table.add_row().cells
        values = (
            f"Synthetic Applicant {number:02d}",
            "PhD",
            first_age if number == 0 else "30",
            first_academic_age if number == 0 else "4",
            "" if number == 0 else "Not stated",
            "1",
            "2",
            "3",
            "4",
            total_citations if number == 0 else "5",
            "NR" if number == 0 else "https://orcid.org/0000-0002-1825-0097",
            scholar_citations,
            "95%",
        )
        for cell, value in zip(cells, values, strict=True):
            cell.text = value
    document.save(path)


def test_parser_finds_the_first_table_with_the_register_header_and_preserves_source_observations(
    tmp_path: Path,
) -> None:
    """Break caught: changing table order or blank source values could mis-import the register."""
    register_path = tmp_path / "register.docx"
    write_register(register_path)

    applicants = parse_register(register_path)

    assert len(applicants) == 36
    first = applicants[0]
    assert first.applicant_name == "Synthetic Applicant 00"
    assert first.degree == "PhD"
    assert first.age_observation == 31
    assert first.academic_age_observation == 4.8
    assert first.gender is None
    assert first.first_author_papers == 1
    assert first.last_author_papers == 2
    assert first.total_papers == 3
    assert first.h_index == 4
    assert first.total_citations == 5
    assert first.total_citations_qualifier is None
    assert first.orcid is None
    assert first.google_scholar_citations is None
    assert first.identity_certainty == "95%"


@pytest.mark.parametrize(
    ("first_age", "first_academic_age", "scholar_citations", "expected"),
    [
        ("31 years", "4.8", "", "age"),
        ("31", "four", "", "academic age"),
        ("31", "4.8", "not-a-count", "Google Scholar"),
    ],
)
def test_parser_rejects_malformed_values_instead_of_coercing_them(
    tmp_path: Path,
    first_age: str,
    first_academic_age: str,
    scholar_citations: str,
    expected: str,
) -> None:
    """Break caught: a malformed register value could silently become a false numeric observation."""
    register_path = tmp_path / "register.docx"
    write_register(
        register_path,
        first_age=first_age,
        first_academic_age=first_academic_age,
        scholar_citations=scholar_citations,
    )

    with pytest.raises(RegisterParseError, match=expected):
        parse_register(register_path)


def test_parser_preserves_a_lower_bound_citation_observation_without_coercing_it(
    tmp_path: Path,
) -> None:
    register_path = tmp_path / "register.docx"
    write_register(register_path, total_citations=">100")

    first = parse_register(register_path)[0]

    assert first.total_citations is None
    assert first.total_citations_qualifier == ">100"
